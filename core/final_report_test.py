import os
from docx import Document
from datetime import datetime
from core.llm_utils import call_llm  # ✅ 공통 LLM 호출 유틸
# ✅ llm_utils 내부에서 LLM_URL, TOKEN, MODEL 모두 관리


def save_report_to_word(report_text: str, output_dir: str = "./reports"):
    """생성된 보고서를 Word 파일로 저장"""
    os.makedirs(output_dir, exist_ok=True)

    filename = f"건설사고_재발방지대책보고서_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(output_dir, filename)

    doc = Document()
    doc.add_heading("건설 사고 재발 방지 대책 보고서", level=0)

    for line in report_text.split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    doc.save(filepath)
    print(f"✅ 보고서가 Word 파일로 저장되었습니다: {filepath}")
    return filepath


rag_output = '''
📝 입력 쿼리: 기본 쿼리: "건축 철근콘크리트공사 설치작업 고소작업 추락 위험 안전난간대 안전고리 미흡" 부스팅 쿼리: "건축 철근콘크리트공사 설치작업 고소작업 추락 위험 안전난간대 안전고리 미흡 법규 기준 지침 체크리스트 조항
(중략)
'''


def generate_accident_report(rag_output: str) -> str:
    """RAG 기반 사고 정보를 입력받아 건설 사고 재발 방지 대책 보고서를 생성"""

    system_message = {
        "role": "system",
        "content": """
당신은 건설 안전 및 사고 재발 방지 보고서를 전문적으로 작성하는 전문가입니다.  
입력으로 제공되는 RAG 분석 결과(<chunk>)에는 ‘사고 개요’, ‘위험 요인’, ‘즉시 조치’, ‘관련 규정’이 포함되어 있습니다.  
이 정보를 바탕으로 **Word 기준 약 4페이지 분량(약 1800~2200 단어)**의 정식 보고서를 작성하십시오.

(중략 — 원문 prompt 그대로 유지)
"""
    }

    user_message = {
        "role": "user",
        "content": f"다음은 RAG 분석 결과이다. 이를 토대로 보고서를 작성하라:\n\n{rag_output}"
    }

    try:
        # ✅ llm_utils에서 API 호출 통합
        report_text = call_llm(
            [system_message, user_message],
            temperature=0.3,
            top_p=0.9,
            max_tokens=25000
        )
        return report_text
    except Exception as e:
        print(f"⚠️ 보고서 생성 실패: {e}")
        return "보고서 생성 실패"


def main():
    report = generate_accident_report(rag_output)
    print("===== 건설 사고 재발 방지 대책 보고서 초안 =====\n")
    print(report)
    save_report_to_word(report)


if __name__ == "__main__":
    main()
