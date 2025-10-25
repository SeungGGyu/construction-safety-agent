import sys, os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))
import streamlit as st
from io import BytesIO
from docx import Document
from typing import Dict, Any, List
from main import run_retrieve, run_rewrite, run_generate

# ==========================
# 페이지 설정
# ==========================
st.set_page_config(page_title="🏗️ 건설 사고 대응 보고서 생성 DEMO", layout="wide")
st.markdown("<h1 style='color:#003366;'>🏗️ 건설 사고 대응 보고서 생성 DEMO</h1>", unsafe_allow_html=True)
st.divider()


# ==========================
# 유틸 함수
# ==========================
def to_human_query(incident: Dict[str, str]) -> str:
    """LangGraph에 넘길 질의문 구성"""
    return "\n".join([
        "[사고 속성]",
        f"- 작업프로세스: {incident.get('process','')}",
        f"- 공종(중분류): {incident.get('construct_type','')}",
        f"- 사고객체(중분류): {incident.get('object_type','')}",
        f"- 사고원인: {incident.get('reason','')}",
    ])


def make_word_file(report_text: str) -> BytesIO:
    doc = Document()
    doc.add_heading("건설 사고 재발 방지 대책 보고서", level=1)
    doc.add_paragraph(report_text)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ==========================
# 세션 초기화
# ==========================
if "state" not in st.session_state:
    st.session_state.state = None
if "related_docs" not in st.session_state:
    st.session_state.related_docs = []
if "report" not in st.session_state:
    st.session_state.report = None


# ==========================
# 1단계: 사고 정보 입력
# ==========================
st.markdown("### 🔹 1단계: 사고 정보 입력")
st.info("작업 정보를 입력한 후 '관련 문서 검색' 버튼을 눌러주세요.")

col1, col2 = st.columns(2)
with col1:
    process = st.text_input("작업프로세스", placeholder="예: 거푸집 해체")
    object_type = st.text_input("사고객체(중분류)", placeholder="예: 개구부/개방형 통로")
with col2:
    construct_type = st.text_input("공종(중분류)", placeholder="예: 골조공사")
    reason = st.text_input("사고원인", placeholder="예: 안전난간 미설치, 덮개 미흡")

st.markdown("---")

# ==========================
# 2단계: 관련 문서 검색
# ==========================
st.markdown("### 🔍 2단계: 관련 문서 검색")

if st.button("📂 관련 문서 검색", use_container_width=True):
    incident = {
        "process": process,
        "construct_type": construct_type,
        "object_type": object_type,
        "reason": reason,
    }
    query_text = to_human_query(incident)
    st.session_state.state = run_retrieve(query_text)
    st.session_state.related_docs = st.session_state.state.get("retrieved", [])
    st.session_state.report = None

docs = st.session_state.related_docs

if docs:
    for d in docs:
        with st.expander(f"📄 {d.get('title', '관련 문서')}"):
            st.caption(f"출처: {d.get('source', '알 수 없음')}")
            st.write("**키워드:**", ", ".join(d.get("keywords", [])))
            st.write(d.get("summary", "요약 없음"))

    st.markdown("---")
    st.markdown("### 🧠 3단계: 관련 문서 검토")
    st.write("이 문서들이 사고와 관련이 있습니까?")

    c1, c2 = st.columns(2)
    with c1:
        if st.button("✅ 관련 있음", use_container_width=True):
            st.session_state.state["route"] = "generate"
            st.session_state.report = None
            st.success("문서가 관련 있음으로 판단되었습니다. 이제 보고서를 생성할 수 있습니다.")
    with c2:
        if st.button("❌ 관련 없음", use_container_width=True):
            st.session_state.state["route"] = "rewrite"
            st.session_state.state = run_rewrite(st.session_state.state)
            st.session_state.related_docs = st.session_state.state.get("retrieved", [])
            st.warning("질의가 리라이트되어 다시 검색되었습니다.")


# ==========================
# 3단계: 보고서 생성
# ==========================
if st.session_state.state and st.session_state.state.get("route") == "generate":
    st.markdown("---")
    st.markdown("### 🧾 4단계: 보고서 생성")
    if st.button("📘 보고서 생성", use_container_width=True, type="primary"):
        st.session_state.state = run_generate(st.session_state.state)
        report = st.session_state.state.get("report", "")
        summary = st.session_state.state.get("report_summary", "")
        st.session_state.report = {"summary": summary, "full_report": report}


# ==========================
# 4단계: 보고서 미리보기 및 다운로드
# ==========================
if st.session_state.report:
    st.markdown("---")
    st.markdown("### 💾 5단계: 보고서 미리보기 및 다운로드")
    st.success(st.session_state.report["summary"])
    st.text_area("📄 보고서 본문", st.session_state.report["full_report"], height=300)

    buffer = make_word_file(st.session_state.report["full_report"])
    st.download_button(
        label="💾 Word 파일로 저장",
        data=buffer,
        file_name="건설사고_재발방지_보고서.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )
